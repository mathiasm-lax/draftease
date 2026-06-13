"""
Generate a realistic *form lease* .docx with placeholder tokens, so the redline
engine has something to run against. In production this file would instead be the
landlord's own form lease, uploaded once and tagged with the same {{tokens}}.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def build(path: str = "sample_form_lease.docx"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("COMMERCIAL OFFICE LEASE")
    run.bold = True
    run.font.size = Pt(15)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("{{property_name}} — Suite {{suite}}").italic = True

    doc.add_paragraph(
        "This Lease is made as of {{lease_date}} between {{landlord}} "
        "(“Landlord”) and {{tenant}} (“Tenant”)."
    )

    def heading(text):
        p = doc.add_paragraph()
        p.add_run(text).bold = True

    heading("ARTICLE 1 — PREMISES AND TERM")
    doc.add_paragraph(
        "1.1  Premises. Landlord leases to Tenant approximately {{rentable_sf}} "
        "rentable square feet located on the {{floor}} floor of the Building."
    )
    doc.add_paragraph(
        "1.2  Term. The term of this Lease shall be {{term_months}} months, "
        "commencing on {{commencement_date}} (the “Commencement Date”) "
        "and expiring on {{expiration_date}}, unless sooner terminated."
    )

    heading("ARTICLE 2 — RENT")
    doc.add_paragraph(
        "2.1  Base Rent. Tenant shall pay Base Rent at the rate of {{base_rent_psf}} "
        "per rentable square foot per annum, payable in equal monthly installments "
        "of {{monthly_rent}}."
    )
    doc.add_paragraph(
        "2.2  Escalations. Commencing on the first anniversary of the Commencement "
        "Date, Base Rent shall increase by {{annual_escalation}} per annum."
    )
    doc.add_paragraph(
        "2.3  Free Rent. Provided Tenant is not in default, Base Rent shall be "
        "abated for the first {{free_rent_months}} months of the Term."
    )

    heading("ARTICLE 3 — SECURITY AND IMPROVEMENTS")
    doc.add_paragraph(
        "3.1  Security Deposit. Tenant shall deposit {{security_deposit}} with "
        "Landlord as security for performance of Tenant’s obligations."
    )
    doc.add_paragraph(
        "3.2  Tenant Improvement Allowance. Landlord shall provide an improvement "
        "allowance of {{ti_allowance_psf}} per rentable square foot toward the cost "
        "of Tenant’s initial improvements."
    )

    heading("ARTICLE 4 — OPTIONS")
    doc.add_paragraph(
        "4.1  Renewal Option. Tenant shall have {{renewal_option}}, exercisable on "
        "not less than nine (9) months’ prior written notice."
    )
    doc.add_paragraph(
        "4.2  Permitted Use. The Premises shall be used solely for {{permitted_use}} "
        "and for no other purpose without Landlord’s prior written consent."
    )

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("LANDLORD: {{landlord}}").bold = True
    sig2 = doc.add_paragraph()
    sig2.add_run("TENANT: {{tenant}}").bold = True

    doc.save(path)
    return path


if __name__ == "__main__":
    out = build()
    print("Wrote", out)
