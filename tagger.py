"""
Deterministic tagging assistant for Draftease.

Takes a *clean* lease (.docx with no {{tokens}}), best-effort detects standard
lease terms, and inserts {{tokens}} where the user confirms — turning a plain
lease into a redline-ready template. No AI, no network.

Public API:
    STANDARD_TOKENS              -> ordered [(key, label), ...]
    autodetect(path)             -> {token_key: detected_phrase}
    apply_tags(in_path, mapping, out_path) -> count of tokens inserted
"""
import copy
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from redline_engine import _iter_paragraphs, _run, _sanitize


def _run_with_child(child, rpr):
    """A <w:r> carrying a clone of a non-text inline element (tab/br/drawing/…),
    preserving its run properties so layout is not lost during a rebuild."""
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    r.append(copy.deepcopy(child))
    return r

# Standard lease-term token library
STANDARD_TOKENS = [
    ("landlord", "Landlord"),
    ("tenant", "Tenant"),
    ("guarantor", "Guarantor"),
    ("premises", "Premises / address"),
    ("base_rent", "Base rent"),
    ("rentable_sf", "Rentable square feet"),
    ("lease_term", "Lease term"),
    ("commencement_date", "Commencement date"),
    ("expiration_date", "Expiration date"),
    ("free_rent", "Free rent / abatement"),
    ("security_deposit", "Security deposit"),
    ("ti_allowance", "TI allowance"),
    ("renewal_option", "Renewal option"),
    ("permitted_use", "Permitted use"),
]

_MONTH = ("(?:January|February|March|April|May|June|July|August|September|"
          "October|November|December)\\s+\\d{1,2},?\\s+\\d{4}")

# best-effort detection patterns; group(1) is the phrase to tokenize
_PATTERNS = {
    "tenant": [r'(?:between|and|,|by)\s+([A-Z][A-Za-z0-9 .,&\'-]{2,45}?)\s*\(\s*["“]?Tenant',
               r'([A-Z][A-Za-z0-9 .,&\'-]{2,45}?)\s*\(\s*["“]?Tenant'],
    "landlord": [r'(?:between|and|,|by)\s+([A-Z][A-Za-z0-9 .,&\'-]{2,45}?)\s*\(\s*["“]?(?:Landlord|Lessor)',
                 r'([A-Z][A-Za-z0-9 .,&\'-]{2,45}?)\s*\(\s*["“]?(?:Landlord|Lessor)'],
    "guarantor": [r'(?:between|and|,|by)\s+([A-Z][A-Za-z0-9 .,&\'-]{2,45}?)\s*\(\s*["“]?Guarantor',
                  r'([A-Z][A-Za-z0-9 .,&\'-]{2,45}?)\s*\(\s*["“]?Guarantor'],
    "base_rent": [r'(\$[\d,]+(?:\.\d{2})?)\s*per\s+(?:rentable\s+)?square\s+foot',
                  r'Base Rent[^$]{0,40}?(\$[\d,]+(?:\.\d{2})?)'],
    "rentable_sf": [r'([\d]{1,3}(?:,\d{3})+|\d{3,})\s+rentable\s+square\s+feet'],
    "lease_term": [r'term\s+of\s+([A-Za-z0-9\-]+\s+(?:months|years))',
                   r'(\d+)\s*[- ]\s*month\s+term'],
    "commencement_date": [r'commenc\w*\s+(?:on\s+)?(' + _MONTH + r')'],
    "expiration_date": [r'expir\w*\s+(?:on\s+)?(' + _MONTH + r')',
                        r'ending\s+(?:on\s+)?(' + _MONTH + r')'],
    "security_deposit": [r'[Ss]ecurity [Dd]eposit[^$]{0,40}?(\$[\d,]+(?:\.\d{2})?)'],
    "ti_allowance": [r'(?:improvement allowance|tenant improvement allowance)[^$]{0,40}?(\$[\d,]+(?:\.\d{2})?)'],
    "free_rent": [r'([A-Za-z]+\s*\(\d+\)|\d+)\s+months?\s+of\s+(?:abated|free)'],
    "premises": [r'located at\s+([0-9][^,\n]{3,50},[^,\n]{2,40},\s*[A-Za-z .]+\s*\d{5})'],
    "permitted_use": [r'used\s+(?:solely\s+)?for\s+([a-z][^.;\n]{5,80}?)(?:\s+and\s+for\s+no|[.;])'],
    "renewal_option": [r'(option to renew[^.;\n]{0,80})'],
}


def autodetect(path: str) -> dict:
    doc = Document(path)
    text = "\n".join(p.text for p in _iter_paragraphs(doc))
    out = {}
    for key, pats in _PATTERNS.items():
        for pat in pats:
            m = re.search(pat, text)
            if m:
                out[key] = m.group(1).strip().rstrip(",")
                break
    return out


def apply_tags(in_path: str, mapping: dict, out_path: str) -> int:
    """Replace each confirmed phrase with its {{token}} in the .docx.

    Non-text inline elements (tabs, breaks, drawings) are preserved: the
    paragraph is rebuilt atom-by-atom so a tab between a Data-Sheet label and its
    value no longer blocks tagging. A phrase is only tagged where its characters
    form a contiguous run of text atoms (i.e. no tab/break splits the phrase
    itself), so we never replace across a preserved element.
    """
    reps = sorted([(v.strip(), k) for k, v in mapping.items() if v and v.strip()],
                  key=lambda x: -len(x[0]))
    if not reps:
        Document(in_path).save(out_path)
        _sanitize(out_path)
        return 0

    W_T, W_R, W_RPR = qn("w:t"), qn("w:r"), qn("w:rPr")
    doc = Document(in_path)
    count = 0
    for p in _iter_paragraphs(doc):
        runs = p._p.findall(W_R)
        if not runs:
            continue
        # Build an ordered atom list: ('t', char, rpr) for text, ('e', element, rpr)
        # for tabs/breaks/drawings/etc. `pos2atom` maps each text-char index -> atom index.
        atoms = []
        pos2atom = []
        chars = []
        for r in runs:
            rpr = r.find(W_RPR)
            for child in r:
                if child.tag == W_RPR:
                    continue
                if child.tag == W_T:
                    for ch in (child.text or ""):
                        pos2atom.append(len(atoms))
                        chars.append(ch)
                        atoms.append(("t", ch, rpr))
                else:
                    atoms.append(("e", child, rpr))
        full = "".join(chars)
        if not full:
            continue

        used = [False] * len(full)
        marks = []  # (text_start, text_end, token)
        for phrase, token in reps:
            start = 0
            while True:
                j = full.find(phrase, start)
                if j < 0:
                    break
                k = j + len(phrase)
                if not any(used[j:k]):
                    a0, a1 = pos2atom[j], pos2atom[k - 1]
                    # contiguous text atoms only (no preserved element splits the phrase)
                    if a1 - a0 == k - 1 - j and all(atoms[a][0] == "t" for a in range(a0, a1 + 1)):
                        marks.append((j, k, token))
                        for x in range(j, k):
                            used[x] = True
                start = k
        if not marks:
            continue
        marks.sort()
        start_token = {s: tok for (s, e, tok) in marks}
        end_of = {s: e for (s, e, tok) in marks}

        new_nodes = []
        buf, buf_rpr = [], None

        def flush():
            nonlocal buf, buf_rpr
            if buf:
                new_nodes.append(_run("".join(buf), buf_rpr))
                buf = []

        tpos = 0
        skip_until = -1
        for kind, payload, rpr in atoms:
            if kind == "t":
                if tpos < skip_until:
                    tpos += 1
                    continue
                if tpos in start_token:
                    flush()
                    new_nodes.append(_run("{{" + start_token[tpos] + "}}", rpr))
                    count += 1
                    skip_until = end_of[tpos]
                    tpos += 1
                    continue
                if buf and rpr is not buf_rpr:
                    flush()
                if not buf:
                    buf_rpr = rpr
                buf.append(payload)
                tpos += 1
            else:
                flush()
                new_nodes.append(_run_with_child(payload, rpr))
        flush()

        at = list(p._p).index(runs[0])
        for n in new_nodes:
            p._p.insert(at, n); at += 1
        for r in runs:
            p._p.remove(r)

    doc.save(out_path)
    _sanitize(out_path)
    return count
