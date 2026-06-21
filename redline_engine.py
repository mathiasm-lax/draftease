"""
Draftease redline engine
=========================

The deterministic core of Draftease. Given:
  * a *form lease* .docx containing placeholder tokens like ``{{base_rent}}``, and
  * a dict of *deal terms* (extracted from a signed LOI) mapping token -> value,

it produces a new .docx in which every placeholder has been replaced with the
deal-specific value **as a native Word tracked change** (a ``w:del`` of the
placeholder followed by a ``w:ins`` of the value). Opening the output in
Microsoft Word shows a real redline that a reviewer can Accept or Reject.

Design notes
------------
* No AI, no network. This is plain OOXML manipulation, so the lease template and
  the resulting redline never leave the machine this runs on.
* Run formatting (``w:rPr``) is preserved: each replacement copies the formatting
  of the text it replaces, so bold / font / size survive.
* Tokens that span multiple runs are handled (Word frequently splits text across
  runs); the paragraph text is reconstructed character-by-character with a map
  back to each character's formatting.
* Paragraphs whose tokenised runs also contain tabs/breaks/drawings are left
  untouched by the rebuild path to avoid dropping those elements (rare in lease
  bodies; logged in the report).

Public API
----------
    generate_redline(template_path, terms, output_path, author="Draftease", date=None) -> dict
"""

from __future__ import annotations

import copy
import datetime as _dt
import re
from typing import Dict

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# {{ token_name }} — letters, digits and underscores, optional surrounding spaces
TOKEN_RE = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")

_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
# run-level child tags that we are not willing to silently drop during a rebuild
_NON_TEXT_RUN_CHILDREN = {qn("w:tab"), qn("w:br"), qn("w:drawing"), qn("w:cr"), qn("w:pict")}


# --------------------------------------------------------------------------- #
# low-level OOXML builders
# --------------------------------------------------------------------------- #
def _clone_rpr(rpr):
    return copy.deepcopy(rpr) if rpr is not None else None


def _run(text: str, rpr):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(_clone_rpr(rpr))
    t = OxmlElement("w:t")
    t.set(_XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    return r


def _del(text: str, rpr, wid: int, author: str, date: str):
    d = OxmlElement("w:del")
    d.set(qn("w:id"), str(wid))
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(_clone_rpr(rpr))
    dt = OxmlElement("w:delText")
    dt.set(_XML_SPACE, "preserve")
    dt.text = text
    r.append(dt)
    d.append(r)
    return d


def _ins(text: str, rpr, wid: int, author: str, date: str):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(wid))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    ins.append(_run(text, rpr))
    return ins


# --------------------------------------------------------------------------- #
# paragraph processing
# --------------------------------------------------------------------------- #
def _direct_runs(p):
    return p.findall(qn("w:r"))


def _paragraph_is_rebuildable(runs) -> bool:
    """True only if every run contains nothing but w:rPr and w:t (safe to rebuild)."""
    for r in runs:
        for child in r:
            if child.tag in _NON_TEXT_RUN_CHILDREN:
                return False
    return True


def _run_with_child(child, rpr):
    """A <w:r> carrying a clone of a non-text inline element (tab/br/drawing/…),
    so layout elements survive a paragraph rebuild."""
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(_clone_rpr(rpr))
    r.append(copy.deepcopy(child))
    return r


def _process_paragraph(p, terms: Dict[str, str], counter, author, date, report) -> None:
    runs = _direct_runs(p)
    if not runs:
        return

    # Build an ordered atom list, preserving non-text inline children (tabs,
    # breaks, drawings) so we can rebuild the paragraph without dropping them.
    W_T, W_RPR = qn("w:t"), qn("w:rPr")
    atoms = []        # ('t', char, rpr) | ('e', element, rpr)
    pos2atom = []     # text-char index -> atom index
    chars: list[str] = []
    for r in runs:
        rpr = r.find(W_RPR)
        for child in r:
            if child.tag == W_RPR:
                continue
            if child.tag == W_T:
                for ch in (child.text or ""):
                    pos2atom.append(len(atoms))
                    chars.append(ch)
                    atoms.append(("t", ch, rpr))
            else:
                atoms.append(("e", child, rpr))
    full = "".join(chars)

    matches = list(TOKEN_RE.finditer(full))
    if not matches:
        return

    # A token is replaceable only if its characters are a contiguous run of text
    # atoms (i.e. no tab/break splits the placeholder itself).
    plan = {}  # text_start -> (text_end, name, raw)
    for m in matches:
        j, k = m.start(), m.end()
        a0, a1 = pos2atom[j], pos2atom[k - 1]
        if a1 - a0 == k - 1 - j and all(atoms[a][0] == "t" for a in range(a0, a1 + 1)):
            plan[j] = (k, m.group(1), m.group(0))
        else:
            report.setdefault("skipped", []).append(m.group(1))
    if not plan:
        return

    new_nodes = []
    buf: list[str] = []
    buf_rpr = None

    def flush():
        nonlocal buf, buf_rpr
        if buf:
            new_nodes.append(_run("".join(buf), buf_rpr))
            buf = []

    tpos = 0
    skip_until = -1
    for kind, payload, rpr in atoms:
        if kind == "t":
            if tpos < skip_until:
                tpos += 1
                continue
            if tpos in plan:
                flush()
                end, name, raw = plan[tpos]
                if name in terms:
                    value = str(terms[name])
                    counter[0] += 1
                    new_nodes.append(_del(raw, rpr, counter[0], author, date))
                    counter[0] += 1
                    new_nodes.append(_ins(value, rpr, counter[0], author, date))
                    report.setdefault("applied", []).append({"token": name, "value": value})
                else:
                    new_nodes.append(_run(raw, rpr))
                    report.setdefault("unmatched", []).append(name)
                skip_until = end
                tpos += 1
                continue
            if buf and rpr is not buf_rpr:
                flush()
            if not buf:
                buf_rpr = rpr
            buf.append(payload)
            tpos += 1
        else:
            flush()
            new_nodes.append(_run_with_child(payload, rpr))
    flush()

    # Splice: insert the new nodes where the first run was, then drop old runs.
    insert_at = list(p).index(runs[0])
    for node in new_nodes:
        p.insert(insert_at, node)
        insert_at += 1
    for r in runs:
        p.remove(r)


# --------------------------------------------------------------------------- #
# document traversal
# --------------------------------------------------------------------------- #
def _iter_paragraphs(document):
    """Yield every <w:p> in the document, wherever it lives.

    Unlike python-docx's structured navigation (which only reaches body
    paragraphs and tables that are *direct* children of the body), this walks
    the raw XML tree, so it also reaches paragraphs nested inside content
    controls (``w:sdt``), tables wrapped in content controls, text boxes, and
    arbitrarily nested tables — common in real lease "Data Sheet" forms.
    Headers and footers (separate XML parts) are covered too.
    """
    from docx.text.paragraph import Paragraph

    P = qn("w:p")
    seen = set()

    def emit(root):
        for p_elem in root.iter(P):
            key = id(p_elem)
            if key in seen:
                continue
            seen.add(key)
            yield Paragraph(p_elem, None)

    yield from emit(document.element.body)
    for section in document.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is None:
                continue
            yield from emit(hf._element)


# --------------------------------------------------------------------------- #
# public API
# --------------------------------------------------------------------------- #
def generate_redline(template_path: str,
                     terms: Dict[str, str],
                     output_path: str,
                     author: str = "Draftease",
                     date: str | None = None) -> dict:
    """
    Fill the placeholders in ``template_path`` with ``terms`` and write a
    tracked-changes .docx to ``output_path``.

    Returns a report dict: {"applied": [...], "unmatched": [...], "skipped": [...]}.
    """
    if date is None:
        date = _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    document = Document(template_path)
    counter = [0]                 # mutable tracked-change id counter
    report: dict = {"applied": [], "unmatched": [], "skipped": []}

    for p in _iter_paragraphs(document):
        _process_paragraph(p._p, terms, counter, author, date, report)

    document.save(output_path)
    _sanitize(output_path)
    return report


def _sanitize(path: str) -> None:
    """Repair a benign python-docx artifact: a <w:zoom/> with no w:percent fails
    strict OOXML validation. Rewrite settings.xml to drop the empty element."""
    import os
    import re as _re
    import shutil
    import zipfile

    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                txt = data.decode("utf-8")
                txt = _re.sub(r"<w:zoom(?![^>]*w:percent)[^>]*/>", "", txt)
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, path)


def generate_redline_direct(template_path: str, changes, output_path: str,
                            author: str = "Draftease", date: str | None = None) -> dict:
    """Redline a lease *without* tokens: each (current_phrase -> new_value) becomes a
    tracked deletion of the current text + insertion of the new value."""
    if date is None:
        date = _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    document = Document(template_path)
    counter = [0]
    applied = []
    reps = sorted([((c or "").strip(), (n or "")) for c, n in changes
                   if (c or "").strip() and (n or "") != "" and (c or "").strip() != (n or "")],
                  key=lambda x: -len(x[0]))
    W_T, W_RPR = qn("w:t"), qn("w:rPr")
    for p in _iter_paragraphs(document):
        el = p._p
        runs = el.findall(qn("w:r"))
        if not runs:
            continue
        atoms = []        # ('t', char, rpr) | ('e', element, rpr)
        pos2atom = []
        chars = []
        for r in runs:
            rpr = r.find(W_RPR)
            for child in r:
                if child.tag == W_RPR:
                    continue
                if child.tag == W_T:
                    for ch in (child.text or ""):
                        pos2atom.append(len(atoms))
                        chars.append(ch)
                        atoms.append(("t", ch, rpr))
                else:
                    atoms.append(("e", child, rpr))
        full = "".join(chars)
        if not full:
            continue
        used = [False] * len(full)
        marks = {}  # text_start -> (text_end, cur, new)
        for cur, new in reps:
            start = 0
            while True:
                j = full.find(cur, start)
                if j < 0:
                    break
                k = j + len(cur)
                if not any(used[j:k]):
                    a0, a1 = pos2atom[j], pos2atom[k - 1]
                    if a1 - a0 == k - 1 - j and all(atoms[a][0] == "t" for a in range(a0, a1 + 1)):
                        marks[j] = (k, cur, new)
                        for x in range(j, k):
                            used[x] = True
                start = k
        if not marks:
            continue
        nodes = []
        buf, buf_rpr = [], None

        def flush():
            nonlocal buf, buf_rpr
            if buf:
                nodes.append(_run("".join(buf), buf_rpr))
                buf = []

        tpos = 0
        skip_until = -1
        for kind, payload, rpr in atoms:
            if kind == "t":
                if tpos < skip_until:
                    tpos += 1
                    continue
                if tpos in marks:
                    flush()
                    end, cur, newv = marks[tpos]
                    counter[0] += 1
                    nodes.append(_del(cur, rpr, counter[0], author, date))
                    counter[0] += 1
                    nodes.append(_ins(newv, rpr, counter[0], author, date))
                    applied.append({"from": cur, "to": newv})
                    skip_until = end
                    tpos += 1
                    continue
                if buf and rpr is not buf_rpr:
                    flush()
                if not buf:
                    buf_rpr = rpr
                buf.append(payload)
                tpos += 1
            else:
                flush()
                nodes.append(_run_with_child(payload, rpr))
        flush()
        at = list(el).index(runs[0])
        for n in nodes:
            el.insert(at, n); at += 1
        for r in runs:
            el.remove(r)

    document.save(output_path)
    _sanitize(output_path)
    return {"applied": applied}


def extract_tokens(template_path: str) -> list:
    """Return the ordered, unique list of {{token}} names found in a .docx."""
    document = Document(template_path)
    seen, out = set(), []
    for p in _iter_paragraphs(document):
        for m in TOKEN_RE.finditer(p.text or ""):
            name = m.group(1)
            if name not in seen:
                seen.add(name)
                out.append(name)
    return out


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) != 4:
        print("usage: python redline_engine.py <template.docx> <terms.json> <output.docx>")
        raise SystemExit(1)

    with open(sys.argv[2]) as fh:
        deal_terms = json.load(fh)
    rep = generate_redline(sys.argv[1], deal_terms, sys.argv[3])
    print(f"Applied {len(rep['applied'])} tracked changes -> {sys.argv[3]}")
    if rep["unmatched"]:
        print("Unmatched tokens (no term supplied):", sorted(set(rep["unmatched"])))
    if rep["skipped"]:
        print("Skipped tokens (complex run):", sorted(set(rep["skipped"])))
